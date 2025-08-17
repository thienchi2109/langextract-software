"""
Document ingestion system for processing various file formats.

This module provides the Ingestor class that handles format detection and text extraction
from PDF, DOCX, and Excel files with appropriate error handling and fallback mechanisms.
"""

import os
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
import mimetypes

# Document processing libraries
import fitz  # PyMuPDF
from docx import Document
from openpyxl import load_workbook
import pandas as pd

from .models import ProcessorInterface
from .exceptions import (
    FileAccessError, 
    ValidationError, 
    LangExtractorError,
    ErrorCategory,
    ErrorSeverity,
    handle_error
)
from .ocr_engine import OCREngine


logger = logging.getLogger(__name__)


class DocumentMetadata:
    """Metadata extracted from documents during processing."""
    
    def __init__(
        self,
        file_path: str,
        file_format: str,
        file_size: int,
        page_count: Optional[int] = None,
        has_images: bool = False,
        extraction_method: str = "direct"
    ):
        self.file_path = file_path
        self.file_format = file_format
        self.file_size = file_size
        self.page_count = page_count
        self.has_images = has_images
        self.extraction_method = extraction_method
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary."""
        return {
            'file_path': self.file_path,
            'file_format': self.file_format,
            'file_size': self.file_size,
            'page_count': self.page_count,
            'has_images': self.has_images,
            'extraction_method': self.extraction_method
        }


class Ingestor(ProcessorInterface):
    """
    Document ingestion system that handles multiple file formats.
    
    Supports PDF, DOCX, and Excel files with format detection and appropriate
    text extraction methods for each format.
    """
    
    # Supported file extensions and their MIME types
    SUPPORTED_FORMATS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel',
        '.csv': 'text/csv'
    }
    
    def __init__(self, ocr_enabled: bool = True, ocr_dpi: int = 350):
        """
        Initialize the Ingestor.
        
        Args:
            ocr_enabled: Enable OCR processing for scanned PDFs (default: True)
            ocr_dpi: DPI setting for OCR processing (default: 350, minimum: 300)
        """
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
        self.ocr_enabled = ocr_enabled
        self._ocr_engine = None
        self.ocr_dpi = max(ocr_dpi, 300)  # Enforce minimum 300 DPI
        
    def supports_format(self, file_path: str) -> bool:
        """
        Check if the file format is supported.
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            True if format is supported, False otherwise
        """
        try:
            file_ext = Path(file_path).suffix.lower()
            return file_ext in self.SUPPORTED_FORMATS
        except Exception as e:
            self.logger.warning(f"Error checking file format for {file_path}: {e}")
            return False
    
    def detect_format(self, file_path: str) -> str:
        """
        Detect the file format based on extension and MIME type.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Detected format string
            
        Raises:
            FileAccessError: If file cannot be accessed or format is unsupported
        """
        try:
            if not os.path.exists(file_path):
                raise FileAccessError(
                    f"File not found: {file_path}",
                    file_path=file_path
                )
            
            file_ext = Path(file_path).suffix.lower()
            
            if not file_ext:
                raise ValidationError(
                    f"File has no extension: {file_path}",
                    field_name="file_extension",
                    field_value=file_ext
                )
            
            if file_ext not in self.SUPPORTED_FORMATS:
                supported_exts = ', '.join(self.SUPPORTED_FORMATS.keys())
                raise ValidationError(
                    f"Unsupported file format: {file_ext}. Supported formats: {supported_exts}",
                    field_name="file_format",
                    field_value=file_ext
                )
            
            # Verify MIME type if possible
            mime_type, _ = mimetypes.guess_type(file_path)
            expected_mime = self.SUPPORTED_FORMATS[file_ext]
            
            if mime_type and mime_type != expected_mime:
                self.logger.warning(
                    f"MIME type mismatch for {file_path}: expected {expected_mime}, got {mime_type}"
                )
            
            return file_ext.lstrip('.')  # Return format without dot
            
        except (FileAccessError, ValidationError):
            raise
        except Exception as e:
            raise handle_error(
                e, 
                self.logger, 
                context={'file_path': file_path, 'operation': 'format_detection'}
            )
    
    def _get_ocr_engine(self) -> OCREngine:
        """
        Get or initialize OCR engine with lazy loading.
        
        Returns:
            Initialized OCR engine instance
        """
        if self._ocr_engine is None:
            self._ocr_engine = OCREngine(dpi=self.ocr_dpi)
        return self._ocr_engine
    
    def process(self, file_path: str) -> str:
        """
        Process a file and extract text content.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Extracted text content
            
        Raises:
            FileAccessError: If file cannot be accessed
            ValidationError: If file format is unsupported
            LangExtractorError: For other processing errors
        """
        try:
            file_format = self.detect_format(file_path)
            
            self.logger.info(f"Processing {file_format} file: {file_path}")
            
            if file_format == 'pdf':
                return self.process_pdf(file_path)
            elif file_format in ['docx', 'doc']:
                return self.process_docx(file_path)
            elif file_format in ['xlsx', 'xls']:
                return self.process_excel(file_path)
            elif file_format == 'csv':
                return self.process_csv(file_path)
            else:
                raise ValidationError(
                    f"No processor available for format: {file_format}",
                    field_name="file_format",
                    field_value=file_format
                )
                
        except (FileAccessError, ValidationError, LangExtractorError):
            raise
        except Exception as e:
            raise handle_error(
                e,
                self.logger,
                context={'file_path': file_path, 'operation': 'file_processing'}
            )
    
    def process_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF files using PyMuPDF and OCR fallback.
        
        Uses direct text extraction first, falls back to OCR for scanned pages.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            FileAccessError: If PDF cannot be opened
            LangExtractorError: For PDF processing errors
        """
        try:
            # If OCR is enabled, use the OCR engine which handles both direct text and OCR
            if self.ocr_enabled:
                ocr_engine = self._get_ocr_engine()
                return ocr_engine.extract_text_from_pdf(file_path, ocr_enabled=True)
            
            # Fallback to direct text extraction only
            return self._extract_text_direct_only(file_path)
            
        except fitz.FileDataError as e:
            raise FileAccessError(
                f"Cannot open PDF file (corrupted or invalid): {file_path}",
                file_path=file_path,
                original_error=e
            )
        except fitz.FileNotFoundError as e:
            raise FileAccessError(
                f"PDF file not found: {file_path}",
                file_path=file_path,
                original_error=e
            )
        except Exception as e:
            raise handle_error(
                e,
                self.logger,
                context={'file_path': file_path, 'operation': 'pdf_processing'}
            )
    
    def _extract_text_direct_only(self, file_path: str) -> str:
        """
        Extract text from PDF using only direct text extraction (no OCR).
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        doc = fitz.open(file_path)
        text_content = []
        has_images = False
        page_count = len(doc)
        
        for page_num in range(page_count):
            page = doc.load_page(page_num)
            
            # Extract text
            page_text = page.get_text()
            
            # Check if page has images (potential scanned content)
            image_list = page.get_images()
            if image_list:
                has_images = True
                self.logger.debug(f"Page {page_num + 1} contains {len(image_list)} images")
            
            # If no text but has images, note that OCR is needed
            if not page_text.strip() and image_list:
                self.logger.warning(
                    f"Page {page_num + 1} has no extractable text but contains images. "
                    "OCR processing may be needed."
                )
                page_text = f"[Page {page_num + 1}: Contains images but no extractable text - OCR required]"
            
            if page_text.strip():
                text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
        
        doc.close()
        
        if not text_content:
            if has_images:
                return "[PDF contains only images - OCR processing required]"
            else:
                return "[PDF contains no extractable text]"
        
        full_text = "\n\n".join(text_content)
        self.logger.info(f"Extracted {len(full_text)} characters from PDF with {page_count} pages")
        
        return full_text
    
    def process_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX files using python-docx.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content including tables
            
        Raises:
            FileAccessError: If DOCX cannot be opened
            LangExtractorError: For DOCX processing errors
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_text = [f"--- Table {table_num + 1} ---"]
                
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if len(table_text) > 1:  # More than just the header
                    text_content.extend(table_text)
            
            if not text_content:
                return "[DOCX file contains no extractable text]"
            
            full_text = "\n\n".join(text_content)
            self.logger.info(
                f"Extracted {len(full_text)} characters from DOCX with "
                f"{len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables"
            )
            
            return full_text
            
        except Exception as e:
            if "not a valid" in str(e).lower() or "corrupted" in str(e).lower():
                raise FileAccessError(
                    f"Cannot open DOCX file (corrupted or invalid): {file_path}",
                    file_path=file_path,
                    original_error=e
                )
            else:
                raise handle_error(
                    e,
                    self.logger,
                    context={'file_path': file_path, 'operation': 'docx_processing'}
                )
    
    def process_excel(self, file_path: str) -> str:
        """
        Extract text from Excel files by converting tables to natural language.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Natural language representation of Excel data
            
        Raises:
            FileAccessError: If Excel file cannot be opened
            LangExtractorError: For Excel processing errors
        """
        try:
            # Try to read with openpyxl first (for .xlsx)
            if file_path.lower().endswith('.xlsx'):
                return self._process_excel_openpyxl(file_path)
            else:
                # Use pandas for .xls files
                return self._process_excel_pandas(file_path)
                
        except Exception as e:
            if "not supported" in str(e).lower() or "corrupted" in str(e).lower():
                raise FileAccessError(
                    f"Cannot open Excel file (corrupted or unsupported): {file_path}",
                    file_path=file_path,
                    original_error=e
                )
            else:
                raise handle_error(
                    e,
                    self.logger,
                    context={'file_path': file_path, 'operation': 'excel_processing'}
                )
    
    def _process_excel_openpyxl(self, file_path: str) -> str:
        """Process Excel file using openpyxl."""
        workbook = load_workbook(file_path, read_only=True, data_only=True)
        text_content = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_text = [f"--- Sheet: {sheet_name} ---"]
            
            # Get all rows with data
            rows_with_data = []
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None and str(cell).strip() for cell in row):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    rows_with_data.append(row_data)
            
            if rows_with_data:
                # Convert to natural language format
                if len(rows_with_data) > 1:
                    # Assume first row is headers
                    headers = rows_with_data[0]
                    data_rows = rows_with_data[1:]
                    
                    sheet_text.append(f"This sheet contains {len(data_rows)} rows of data with columns: {', '.join(headers)}")
                    
                    # Add sample rows for context
                    sample_size = min(5, len(data_rows))
                    for i, row in enumerate(data_rows[:sample_size]):
                        row_desc = []
                        for header, value in zip(headers, row):
                            if value and str(value).strip():
                                row_desc.append(f"{header}: {value}")
                        
                        if row_desc:
                            sheet_text.append(f"Row {i + 1}: {', '.join(row_desc)}")
                    
                    if len(data_rows) > sample_size:
                        sheet_text.append(f"... and {len(data_rows) - sample_size} more rows")
                else:
                    # Single row, treat as headers or labels
                    sheet_text.append(f"Contains: {', '.join(rows_with_data[0])}")
            
            if len(sheet_text) > 1:  # More than just the header
                text_content.extend(sheet_text)
        
        workbook.close()
        
        if not text_content:
            return "[Excel file contains no extractable data]"
        
        full_text = "\n\n".join(text_content)
        self.logger.info(f"Extracted {len(full_text)} characters from Excel with {len(workbook.sheetnames)} sheets")
        
        return full_text
    
    def _process_excel_pandas(self, file_path: str) -> str:
        """Process Excel file using pandas (for .xls files)."""
        # Read all sheets
        excel_data = pd.read_excel(file_path, sheet_name=None, engine='xlrd')
        text_content = []
        
        for sheet_name, df in excel_data.items():
            if df.empty:
                continue
                
            sheet_text = [f"--- Sheet: {sheet_name} ---"]
            
            # Get basic info
            rows, cols = df.shape
            sheet_text.append(f"This sheet contains {rows} rows and {cols} columns: {', '.join(df.columns)}")
            
            # Add sample data
            sample_size = min(5, len(df))
            for i, (_, row) in enumerate(df.head(sample_size).iterrows()):
                row_desc = []
                for col, value in row.items():
                    if pd.notna(value) and str(value).strip():
                        row_desc.append(f"{col}: {value}")
                
                if row_desc:
                    sheet_text.append(f"Row {i + 1}: {', '.join(row_desc)}")
            
            if len(df) > sample_size:
                sheet_text.append(f"... and {len(df) - sample_size} more rows")
            
            text_content.extend(sheet_text)
        
        if not text_content:
            return "[Excel file contains no extractable data]"
        
        full_text = "\n\n".join(text_content)
        self.logger.info(f"Extracted {len(full_text)} characters from Excel with {len(excel_data)} sheets")
        
        return full_text
    
    def process_csv(self, file_path: str) -> str:
        """
        Extract text from CSV files.
        
        Args:
            file_path: Path to the CSV file
            
        Returns:
            Natural language representation of CSV data
            
        Raises:
            FileAccessError: If CSV file cannot be opened
            LangExtractorError: For CSV processing errors
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except (UnicodeDecodeError, pd.errors.EmptyDataError):
                    continue
            
            if df is None:
                # Check if file is actually empty
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                        if not content:
                            return "[CSV file contains no data]"
                except:
                    pass
                
                raise FileAccessError(
                    f"Cannot read CSV file with any supported encoding: {file_path}",
                    file_path=file_path
                )
            
            if df.empty:
                return "[CSV file contains no data]"
            
            text_content = []
            rows, cols = df.shape
            text_content.append(f"CSV file contains {rows} rows and {cols} columns: {', '.join(df.columns)}")
            
            # Add sample data
            sample_size = min(5, len(df))
            for i, (_, row) in enumerate(df.head(sample_size).iterrows()):
                row_desc = []
                for col, value in row.items():
                    if pd.notna(value) and str(value).strip():
                        row_desc.append(f"{col}: {value}")
                
                if row_desc:
                    text_content.append(f"Row {i + 1}: {', '.join(row_desc)}")
            
            if len(df) > sample_size:
                text_content.append(f"... and {len(df) - sample_size} more rows")
            
            full_text = "\n\n".join(text_content)
            self.logger.info(f"Extracted {len(full_text)} characters from CSV with {rows} rows")
            
            return full_text
            
        except Exception as e:
            raise handle_error(
                e,
                self.logger,
                context={'file_path': file_path, 'operation': 'csv_processing'}
            )
    
    def get_metadata(self, file_path: str) -> DocumentMetadata:
        """
        Get metadata for a processed document.
        
        Args:
            file_path: Path to the file
            
        Returns:
            DocumentMetadata object with file information
            
        Raises:
            FileAccessError: If file cannot be accessed
        """
        try:
            if not os.path.exists(file_path):
                raise FileAccessError(
                    f"File not found: {file_path}",
                    file_path=file_path
                )
            
            file_format = self.detect_format(file_path)
            file_size = os.path.getsize(file_path)
            page_count = None
            has_images = False
            
            # Get format-specific metadata
            if file_format == 'pdf':
                try:
                    doc = fitz.open(file_path)
                    page_count = len(doc)
                    
                    # Check for images
                    for page_num in range(min(3, page_count)):  # Check first 3 pages
                        page = doc.load_page(page_num)
                        if page.get_images():
                            has_images = True
                            break
                    
                    doc.close()
                except Exception as e:
                    self.logger.warning(f"Could not extract PDF metadata: {e}")
            
            elif file_format in ['xlsx', 'xls']:
                try:
                    if file_format == 'xlsx':
                        workbook = load_workbook(file_path, read_only=True)
                        page_count = len(workbook.sheetnames)
                        workbook.close()
                    else:
                        excel_data = pd.read_excel(file_path, sheet_name=None, engine='xlrd')
                        page_count = len(excel_data)
                except Exception as e:
                    self.logger.warning(f"Could not extract Excel metadata: {e}")
            
            return DocumentMetadata(
                file_path=file_path,
                file_format=file_format,
                file_size=file_size,
                page_count=page_count,
                has_images=has_images,
                extraction_method="direct"
            )
            
        except (FileAccessError, ValidationError):
            raise
        except Exception as e:
            raise handle_error(
                e,
                self.logger,
                context={'file_path': file_path, 'operation': 'metadata_extraction'}
            )
    
    def batch_process(self, file_paths: List[str]) -> Dict[str, Tuple[str, Optional[Exception]]]:
        """
        Process multiple files and return results.
        
        Args:
            file_paths: List of file paths to process
            
        Returns:
            Dictionary mapping file paths to (text_content, error) tuples
        """
        results = {}
        
        for file_path in file_paths:
            try:
                text_content = self.process(file_path)
                results[file_path] = (text_content, None)
                self.logger.info(f"Successfully processed: {file_path}")
            except Exception as e:
                results[file_path] = ("", e)
                self.logger.error(f"Failed to process {file_path}: {e}")
        
        return results
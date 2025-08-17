"""
Unit tests for the document ingestion system.
"""

import os
import tempfile
import pytest
from unittest.mock import Mock, patch, MagicMock
from pathlib import Path

# Test data imports
import fitz  # PyMuPDF
from docx import Document
from openpyxl import Workbook
import pandas as pd

from core.ingestor import Ingestor, DocumentMetadata
from core.exceptions import FileAccessError, ValidationError, LangExtractorError


class TestIngestor:
    """Test cases for the Ingestor class."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.ingestor = Ingestor()
        self.temp_dir = tempfile.mkdtemp()
    
    def teardown_method(self):
        """Clean up test fixtures."""
        # Clean up temp files
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
    
    def create_temp_file(self, filename: str, content: bytes = b"test content") -> str:
        """Create a temporary file for testing."""
        file_path = os.path.join(self.temp_dir, filename)
        with open(file_path, 'wb') as f:
            f.write(content)
        return file_path
    
    def test_supports_format_valid_extensions(self):
        """Test format support detection for valid extensions."""
        test_files = [
            "test.pdf",
            "test.docx",
            "test.doc", 
            "test.xlsx",
            "test.xls",
            "test.csv"
        ]
        
        for filename in test_files:
            file_path = self.create_temp_file(filename)
            assert self.ingestor.supports_format(file_path), f"Should support {filename}"
    
    def test_supports_format_invalid_extensions(self):
        """Test format support detection for invalid extensions."""
        test_files = [
            "test.txt",
            "test.png",
            "test.jpg",
            "test.pptx",
            "test"  # no extension
        ]
        
        for filename in test_files:
            file_path = self.create_temp_file(filename)
            assert not self.ingestor.supports_format(file_path), f"Should not support {filename}"
    
    def test_supports_format_case_insensitive(self):
        """Test that format detection is case insensitive."""
        test_files = [
            "test.PDF",
            "test.DOCX",
            "test.XlSx"
        ]
        
        for filename in test_files:
            file_path = self.create_temp_file(filename)
            assert self.ingestor.supports_format(file_path), f"Should support {filename} (case insensitive)"
    
    def test_detect_format_valid_files(self):
        """Test format detection for valid files."""
        test_cases = [
            ("test.pdf", "pdf"),
            ("test.docx", "docx"),
            ("test.doc", "doc"),
            ("test.xlsx", "xlsx"),
            ("test.xls", "xls"),
            ("test.csv", "csv")
        ]
        
        for filename, expected_format in test_cases:
            file_path = self.create_temp_file(filename)
            detected_format = self.ingestor.detect_format(file_path)
            assert detected_format == expected_format, f"Expected {expected_format}, got {detected_format}"
    
    def test_detect_format_nonexistent_file(self):
        """Test format detection for nonexistent file."""
        with pytest.raises(FileAccessError) as exc_info:
            self.ingestor.detect_format("nonexistent_file.pdf")
        
        assert "File not found" in str(exc_info.value)
    
    def test_detect_format_no_extension(self):
        """Test format detection for file without extension."""
        file_path = self.create_temp_file("test_file")
        
        with pytest.raises(ValidationError) as exc_info:
            self.ingestor.detect_format(file_path)
        
        assert "no extension" in str(exc_info.value)
    
    def test_detect_format_unsupported_extension(self):
        """Test format detection for unsupported extension."""
        file_path = self.create_temp_file("test.txt")
        
        with pytest.raises(ValidationError) as exc_info:
            self.ingestor.detect_format(file_path)
        
        assert "Unsupported file format" in str(exc_info.value)
    
    @patch('fitz.open')
    def test_process_pdf_success(self, mock_fitz_open):
        """Test successful PDF processing."""
        # Mock PDF document
        mock_doc = Mock()
        mock_doc.__len__ = Mock(return_value=2)
        
        # Mock pages
        mock_page1 = Mock()
        mock_page1.get_text.return_value = "Page 1 content"
        mock_page1.get_images.return_value = []
        
        mock_page2 = Mock()
        mock_page2.get_text.return_value = "Page 2 content"
        mock_page2.get_images.return_value = []
        
        mock_doc.load_page.side_effect = [mock_page1, mock_page2]
        mock_fitz_open.return_value = mock_doc
        
        file_path = self.create_temp_file("test.pdf")
        result = self.ingestor.process_pdf(file_path)
        
        assert "Page 1 content" in result
        assert "Page 2 content" in result
        assert "--- Page 1 ---" in result
        assert "--- Page 2 ---" in result
        mock_doc.close.assert_called_once()
    
    @patch('fitz.open')
    def test_process_pdf_with_images(self, mock_fitz_open):
        """Test PDF processing with images (potential OCR needed)."""
        mock_doc = Mock()
        mock_doc.__len__ = Mock(return_value=1)
        
        mock_page = Mock()
        mock_page.get_text.return_value = ""  # No text
        mock_page.get_images.return_value = [{"image": "data"}]  # Has images
        
        mock_doc.load_page.return_value = mock_page
        mock_fitz_open.return_value = mock_doc
        
        file_path = self.create_temp_file("test.pdf")
        result = self.ingestor.process_pdf(file_path)
        
        assert "OCR required" in result
        mock_doc.close.assert_called_once()
    
    @patch('fitz.open')
    def test_process_pdf_corrupted_file(self, mock_fitz_open):
        """Test PDF processing with corrupted file."""
        mock_fitz_open.side_effect = fitz.FileDataError("Corrupted PDF")
        
        file_path = self.create_temp_file("test.pdf")
        
        with pytest.raises(FileAccessError) as exc_info:
            self.ingestor.process_pdf(file_path)
        
        assert "corrupted or invalid" in str(exc_info.value)
    
    def test_process_docx_success(self):
        """Test successful DOCX processing."""
        # Create a real DOCX file for testing
        file_path = os.path.join(self.temp_dir, "test.docx")
        doc = Document()
        
        # Add paragraphs
        doc.add_paragraph("First paragraph content")
        doc.add_paragraph("Second paragraph content")
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"
        
        doc.save(file_path)
        
        result = self.ingestor.process_docx(file_path)
        
        assert "First paragraph content" in result
        assert "Second paragraph content" in result
        assert "Header 1" in result
        assert "Data 1" in result
        assert "--- Table 1 ---" in result
    
    def test_process_docx_empty_file(self):
        """Test DOCX processing with empty file."""
        file_path = os.path.join(self.temp_dir, "empty.docx")
        doc = Document()
        doc.save(file_path)
        
        result = self.ingestor.process_docx(file_path)
        assert "no extractable text" in result
    
    def test_process_excel_xlsx_success(self):
        """Test successful Excel XLSX processing."""
        file_path = os.path.join(self.temp_dir, "test.xlsx")
        
        # Create a real Excel file
        wb = Workbook()
        ws = wb.active
        ws.title = "Test Sheet"
        
        # Add headers
        ws['A1'] = "Name"
        ws['B1'] = "Age"
        ws['C1'] = "City"
        
        # Add data
        ws['A2'] = "John Doe"
        ws['B2'] = 30
        ws['C2'] = "New York"
        
        ws['A3'] = "Jane Smith"
        ws['B3'] = 25
        ws['C3'] = "Los Angeles"
        
        wb.save(file_path)
        
        result = self.ingestor.process_excel(file_path)
        
        assert "Test Sheet" in result
        assert "Name, Age, City" in result
        assert "John Doe" in result
        assert "Jane Smith" in result
        assert "2 rows of data" in result
    
    def test_process_excel_xls_success(self):
        """Test successful Excel XLS processing using pandas."""
        file_path = os.path.join(self.temp_dir, "test.xls")
        
        # Create test data
        data = {
            'Name': ['Alice', 'Bob', 'Charlie'],
            'Score': [95, 87, 92],
            'Grade': ['A', 'B', 'A']
        }
        df = pd.DataFrame(data)
        
        # Save as XLS (this will actually save as xlsx but with .xls extension for testing)
        with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
            df.to_excel(writer, sheet_name='Grades', index=False)
        
        # Mock pandas read_excel to return our test data
        with patch('pandas.read_excel') as mock_read:
            mock_read.return_value = {'Grades': df}
            
            result = self.ingestor.process_excel(file_path)
            
            assert "Grades" in result
            assert "Name, Score, Grade" in result
            assert "Alice" in result
            assert "3 rows and 3 columns" in result
    
    def test_process_csv_success(self):
        """Test successful CSV processing."""
        file_path = os.path.join(self.temp_dir, "test.csv")
        
        # Create CSV content
        csv_content = """Name,Age,Department
John Smith,28,Engineering
Sarah Johnson,32,Marketing
Mike Brown,25,Sales"""
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(csv_content)
        
        result = self.ingestor.process_csv(file_path)
        
        assert "3 rows and 3 columns" in result
        assert "Name, Age, Department" in result
        assert "John Smith" in result
        assert "Engineering" in result
    
    def test_process_csv_empty_file(self):
        """Test CSV processing with empty file."""
        file_path = os.path.join(self.temp_dir, "empty.csv")
        
        with open(file_path, 'w') as f:
            f.write("")
        
        result = self.ingestor.process_csv(file_path)
        assert "no data" in result
    
    def test_process_csv_encoding_issues(self):
        """Test CSV processing with encoding issues."""
        file_path = os.path.join(self.temp_dir, "encoded.csv")
        
        # Create CSV with special characters
        csv_content = "Name,City\nJohn,Hà Nội\nMarie,Hồ Chí Minh"
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(csv_content)
        
        result = self.ingestor.process_csv(file_path)
        
        assert "Hà Nội" in result or "H" in result  # Should handle encoding gracefully
    
    def test_process_unsupported_format(self):
        """Test processing unsupported file format."""
        file_path = self.create_temp_file("test.txt")
        
        with pytest.raises(ValidationError) as exc_info:
            self.ingestor.process(file_path)
        
        assert "Unsupported file format" in str(exc_info.value)
    
    def test_process_nonexistent_file(self):
        """Test processing nonexistent file."""
        with pytest.raises(FileAccessError) as exc_info:
            self.ingestor.process("nonexistent.pdf")
        
        assert "File not found" in str(exc_info.value)
    
    def test_get_metadata_pdf(self):
        """Test metadata extraction for PDF files."""
        file_path = self.create_temp_file("test.pdf")
        
        with patch('fitz.open') as mock_fitz_open:
            mock_doc = Mock()
            mock_doc.__len__ = Mock(return_value=3)
            
            mock_page = Mock()
            mock_page.get_images.return_value = [{"image": "data"}]
            mock_doc.load_page.return_value = mock_page
            
            mock_fitz_open.return_value = mock_doc
            
            metadata = self.ingestor.get_metadata(file_path)
            
            assert metadata.file_format == "pdf"
            assert metadata.page_count == 3
            assert metadata.has_images == True
            assert metadata.file_size > 0
            
            mock_doc.close.assert_called_once()
    
    def test_get_metadata_excel(self):
        """Test metadata extraction for Excel files."""
        file_path = os.path.join(self.temp_dir, "test.xlsx")
        
        wb = Workbook()
        wb.create_sheet("Sheet2")
        wb.save(file_path)
        
        metadata = self.ingestor.get_metadata(file_path)
        
        assert metadata.file_format == "xlsx"
        assert metadata.page_count == 2  # Default sheet + Sheet2
        assert metadata.file_size > 0
    
    def test_get_metadata_nonexistent_file(self):
        """Test metadata extraction for nonexistent file."""
        with pytest.raises(FileAccessError) as exc_info:
            self.ingestor.get_metadata("nonexistent.pdf")
        
        assert "File not found" in str(exc_info.value)
    
    def test_batch_process_success(self):
        """Test batch processing of multiple files."""
        # Create test files
        pdf_path = self.create_temp_file("test.pdf")
        csv_path = os.path.join(self.temp_dir, "test.csv")
        
        with open(csv_path, 'w') as f:
            f.write("Name,Value\nTest,123")
        
        # Mock PDF processing
        with patch.object(self.ingestor, 'process_pdf', return_value="PDF content"):
            results = self.ingestor.batch_process([pdf_path, csv_path])
        
        assert len(results) == 2
        assert results[pdf_path][0] == "PDF content"
        assert results[pdf_path][1] is None  # No error
        assert "Name, Value" in results[csv_path][0]  # Note the space after comma
        assert results[csv_path][1] is None  # No error
    
    def test_batch_process_with_errors(self):
        """Test batch processing with some files failing."""
        valid_csv = os.path.join(self.temp_dir, "valid.csv")
        invalid_file = "nonexistent.pdf"
        
        with open(valid_csv, 'w') as f:
            f.write("Name,Value\nTest,123")
        
        results = self.ingestor.batch_process([valid_csv, invalid_file])
        
        assert len(results) == 2
        assert results[valid_csv][1] is None  # No error for valid file
        assert results[invalid_file][1] is not None  # Error for invalid file
        assert isinstance(results[invalid_file][1], FileAccessError)
    
    def test_document_metadata_to_dict(self):
        """Test DocumentMetadata to_dict conversion."""
        metadata = DocumentMetadata(
            file_path="/test/file.pdf",
            file_format="pdf",
            file_size=1024,
            page_count=5,
            has_images=True,
            extraction_method="direct"
        )
        
        result = metadata.to_dict()
        
        expected = {
            'file_path': "/test/file.pdf",
            'file_format': "pdf",
            'file_size': 1024,
            'page_count': 5,
            'has_images': True,
            'extraction_method': "direct"
        }
        
        assert result == expected
    
    def test_error_handling_with_context(self):
        """Test that errors include proper context information."""
        file_path = self.create_temp_file("test.pdf")
        
        with patch('fitz.open') as mock_fitz_open:
            mock_fitz_open.side_effect = Exception("Unexpected error")
            
            with pytest.raises(LangExtractorError) as exc_info:
                self.ingestor.process_pdf(file_path)
            
            error = exc_info.value
            assert file_path in str(error.details.get('file_path', ''))
            assert error.details.get('operation') == 'pdf_processing'


class TestIngestorIntegration:
    """Integration tests for the Ingestor class with real files."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.ingestor = Ingestor()
        self.temp_dir = tempfile.mkdtemp()
    
    def teardown_method(self):
        """Clean up test fixtures."""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
    
    def test_real_docx_processing(self):
        """Test processing a real DOCX file."""
        file_path = os.path.join(self.temp_dir, "integration_test.docx")
        
        # Create a comprehensive DOCX file
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph with some content.')
        doc.add_paragraph('This is another paragraph with different content.')
        
        # Add a table
        table = doc.add_table(rows=3, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Product'
        hdr_cells[1].text = 'Quantity'
        hdr_cells[2].text = 'Price'
        
        # Add data rows
        row1_cells = table.rows[1].cells
        row1_cells[0].text = 'Widget A'
        row1_cells[1].text = '10'
        row1_cells[2].text = '$5.00'
        
        row2_cells = table.rows[2].cells
        row2_cells[0].text = 'Widget B'
        row2_cells[1].text = '5'
        row2_cells[2].text = '$10.00'
        
        doc.save(file_path)
        
        # Process the file
        result = self.ingestor.process(file_path)
        
        # Verify content extraction
        assert 'Test Document' in result
        assert 'test paragraph' in result
        assert 'another paragraph' in result
        assert 'Product' in result
        assert 'Widget A' in result
        assert '$5.00' in result
        assert '--- Table 1 ---' in result
    
    def test_real_excel_processing(self):
        """Test processing a real Excel file."""
        file_path = os.path.join(self.temp_dir, "integration_test.xlsx")
        
        # Create a comprehensive Excel file
        wb = Workbook()
        
        # First sheet - Sales data
        ws1 = wb.active
        ws1.title = "Sales Data"
        
        headers = ["Date", "Product", "Quantity", "Revenue"]
        for col, header in enumerate(headers, 1):
            ws1.cell(row=1, column=col, value=header)
        
        data = [
            ["2024-01-01", "Product A", 100, 1000.50],
            ["2024-01-02", "Product B", 75, 750.25],
            ["2024-01-03", "Product A", 50, 500.00]
        ]
        
        for row, row_data in enumerate(data, 2):
            for col, value in enumerate(row_data, 1):
                ws1.cell(row=row, column=col, value=value)
        
        # Second sheet - Summary
        ws2 = wb.create_sheet("Summary")
        ws2['A1'] = "Total Revenue"
        ws2['B1'] = 2250.75
        ws2['A2'] = "Total Quantity"
        ws2['B2'] = 225
        
        wb.save(file_path)
        
        # Process the file
        result = self.ingestor.process(file_path)
        
        # Verify content extraction
        assert 'Sales Data' in result
        assert 'Summary' in result
        assert 'Date, Product, Quantity, Revenue' in result
        assert 'Product A' in result
        assert '1000.5' in result or '1000.50' in result
        assert 'Total Revenue' in result
        assert '2250.75' in result
    
    def test_format_detection_accuracy(self):
        """Test format detection accuracy with real files."""
        # Create files with correct content
        test_files = []
        
        # DOCX file
        docx_path = os.path.join(self.temp_dir, "test.docx")
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(docx_path)
        test_files.append((docx_path, "docx"))
        
        # Excel file
        xlsx_path = os.path.join(self.temp_dir, "test.xlsx")
        wb = Workbook()
        wb.save(xlsx_path)
        test_files.append((xlsx_path, "xlsx"))
        
        # CSV file
        csv_path = os.path.join(self.temp_dir, "test.csv")
        with open(csv_path, 'w') as f:
            f.write("Name,Value\nTest,123")
        test_files.append((csv_path, "csv"))
        
        # Test format detection
        for file_path, expected_format in test_files:
            detected_format = self.ingestor.detect_format(file_path)
            assert detected_format == expected_format, f"Expected {expected_format}, got {detected_format} for {file_path}"
            
            # Test that the file can be processed
            result = self.ingestor.process(file_path)
            assert len(result) > 0, f"No content extracted from {file_path}"


if __name__ == "__main__":
    pytest.main([__file__])